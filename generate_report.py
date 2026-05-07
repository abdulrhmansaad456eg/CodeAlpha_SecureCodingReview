import subprocess
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

import os
from datetime import datetime

def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_code_block(doc, code_lines, is_vulnerable=False):
    table = doc.add_table(rows=len(code_lines), cols=2)
    table.style = 'Table Grid'
    for i, (line_num, code) in enumerate(code_lines):
        num_cell = table.rows[i].cells[0]
        num_cell.text = str(line_num)
        num_cell.width = Inches(0.5)
        for paragraph in num_cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)
        
        code_cell = table.rows[i].cells[1]
        code_cell.text = code
        code_cell.width = Inches(5.5)
        for paragraph in code_cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                if is_vulnerable:
                    run.font.color.rgb = RGBColor(200, 0, 0)
                else:
                    run.font.color.rgb = RGBColor(0, 100, 0)

def generate_report():
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    title = doc.add_paragraph()
    title_run = title.add_run('Secure Coding Review of a Flask Web Application')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Cybersecurity Internship Project')
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info_run = info.add_run('Prepared by: Cybersecurity Intern\nDate: May 2026\nDuration: 4 weeks')
    info_run.font.size = Pt(12)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    add_heading_custom(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction',
        '2. Project Scope and Objectives',
        '3. Methodology',
        '4. Tools Used',
        '5. Vulnerability Findings',
        '6. Risk Severity Assessment',
        '7. Remediation Steps',
        '8. Secure Coding Best Practices',
        '9. Comparison Analysis',
        '10. Conclusion',
        '11. References'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    add_heading_custom(doc, '1. Introduction', 1)
    doc.add_paragraph(
        'This report documents a comprehensive secure coding review conducted as part of a '
        'cybersecurity internship program. The project involved analyzing a deliberately '
        'vulnerable Flask web application, identifying security flaws through both '
        'automated scanning tools and manual code review, and subsequently implementing '
        'remediation measures to create a secure version of the application.'
    )
    doc.add_paragraph(
        'The primary goal of this exercise was to gain practical experience in secure '
        'software development practices, vulnerability assessment, and secure coding principles '
        'that align with industry standards such as the OWASP Top 10 and CWE Top 25.'
    )
    
    add_heading_custom(doc, '2. Project Scope and Objectives', 1)
    add_heading_custom(doc, '2.1 Project Scope', 2)
    doc.add_paragraph(
        'The scope of this secure coding review encompassed a complete Flask web application '
        'including user authentication, file upload functionality, search capabilities, and '
        'administrative features. The application uses SQLite as the database backend and '
        'follows a typical Model-View-Controller architecture.'
    )
    
    add_heading_custom(doc, '2.2 Project Objectives', 2)
    objectives = [
        'Identify security vulnerabilities in the vulnerable Flask application',
        'Perform static analysis using automated security scanning tools',
        'Conduct manual code review to discover logic flaws',
        'Assess the severity and business impact of identified vulnerabilities',
        'Implement secure remediation measures',
        'Document findings and recommendations in a professional report'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    # 3. Methodology
    add_heading_custom(doc, '3. Methodology', 1)
    doc.add_paragraph(
        'The secure coding review followed a structured methodology combining automated '
        'scanning with manual analysis to ensure comprehensive vulnerability coverage.'
    )
    
    add_heading_custom(doc, '3.1 Automated Static Analysis', 2)
    doc.add_paragraph(
        'Bandit and Semgrep were employed to perform automated security scanning. These tools '
        'analyze source code for common security anti-patterns and vulnerabilities without '
        'requiring the application to be running.'
    )
    
    add_heading_custom(doc, '3.2 Manual Code Review', 2)
    doc.add_paragraph(
        'Manual inspection of the source code was performed to identify business logic flaws, '
        'insecure coding patterns, and vulnerabilities that automated tools might miss. '
        'The review focused on authentication mechanisms, data validation, and session management.'
    )
    
    add_heading_custom(doc, '3.3 Remediation and Validation', 2)
    doc.add_paragraph(
        'After identifying vulnerabilities, a remediated version of the application was '
        'developed. Each fix was validated to ensure the vulnerability was properly addressed '
        'while maintaining application functionality.'
    )
    
    add_heading_custom(doc, '4. Tools Used', 1)
    
    tools_data = [
        ('Tool', 'Version', 'Purpose'),
        ('Bandit', '1.7.5', 'Python-specific security linter detecting common security issues'),
        ('Semgrep', '1.45.0', 'Static analysis engine for identifying code patterns and vulnerabilities'),
        ('Flask', '2.3.3', 'Web framework for the application being analyzed'),
        ('SQLite', '3.x', 'Database engine for data storage'),
        ('Python', '3.10+', 'Programming language for the application')
    ]
    
    tools_table = doc.add_table(rows=len(tools_data), cols=3)
    tools_table.style = 'Table Grid'
    
    for i, (tool, version, purpose) in enumerate(tools_data):
        row = tools_table.rows[i]
        row.cells[0].text = tool
        row.cells[1].text = version
        row.cells[2].text = purpose
        
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    cell._element.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="003366"/>'.format(nsdecls('w')))
                    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '5. Vulnerability Findings', 1)
    doc.add_paragraph(
        'The security assessment identified nine distinct vulnerabilities across the application. '
        'These vulnerabilities span multiple categories including injection flaws, broken '
        'authentication, security misconfiguration, and sensitive data exposure.'
    )
    
    add_heading_custom(doc, '5.1 SQL Injection (CWE-89)', 2)
    p = doc.add_paragraph()
    p.add_run('Severity: ').bold = True
    p.add_run('CRITICAL')
    p.runs[1].font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph(
        'The application constructs SQL queries using string concatenation with user-supplied '
        'input, making it vulnerable to SQL injection attacks. An attacker can manipulate '
        'the input to alter query logic, bypass authentication, or extract sensitive data.'
    )
    
    add_code_block(doc, [
        (86, "query = f\"SELECT * FROM users WHERE username = '{username}'\""),
        (87, "             AND password = '{password}'\""),
        (88, 'cursor.execute(query)')
    ], is_vulnerable=True)
    
    add_code_block(doc, [
        (87, 'cursor.execute('),
        (88, "    'SELECT * FROM users WHERE username = ? AND password = ?',"),
        (89, '    (username, password)'),
        (90, ')')
    ], is_vulnerable=False)
    
    add_heading_custom(doc, '5.2 Hardcoded Secret Key (CWE-798)', 2)
    p = doc.add_paragraph()
    p.add_run('Severity: ').bold = True
    p.add_run('HIGH')
    p.runs[1].font.color.rgb = RGBColor(255, 140, 0)
    
    doc.add_paragraph(
        'The Flask secret key is hardcoded in the application source code. This key is used '
        'for session management and cryptographic operations. Anyone with access to the source '
        'code can forge session cookies and impersonate users.'
    )
    
    add_code_block(doc, [
        (21, "app.secret_key = 'super_secret_key_12345_do_not_change'")
    ], is_vulnerable=True)
    
    add_heading_custom(doc, '5.3 Debug Mode Enabled (CWE-489)', 2)
    p = doc.add_paragraph()
    p.add_run('Severity: ').bold = True
    p.add_run('HIGH')
    p.runs[1].font.color.rgb = RGBColor(255, 140, 0)
    
    doc.add_paragraph(
        'Flask debug mode is enabled in the production configuration. This exposes the '
        'Werkzeug interactive debugger, which allows arbitrary code execution through the '
        'browser interface. Stack traces with sensitive information are also displayed.'
    )
    
    add_heading_custom(doc, '5.4 Plaintext Password Storage (CWE-256)', 2)
    p = doc.add_paragraph()
    p.add_run('Severity: ').bold = True
    p.add_run('HIGH')
    p.runs[1].font.color.rgb = RGBColor(255, 140, 0)
    
    doc.add_paragraph(
        'User passwords are stored in plaintext without any hashing. If the database is '
        'compromised, attackers immediately have access to all user credentials. This violates '
        'fundamental security principles and regulatory compliance requirements.'
    )
    
    add_heading_custom(doc, '5.5 Unsafe File Upload (CWE-434)', 2)
    p = doc.add_paragraph()
    p.add_run('Severity: ').bold = True
    p.add_run('MEDIUM')
    p.runs[1].font.color.rgb = RGBColor(255, 165, 0)
    
    doc.add_paragraph(
        'The file upload functionality does not validate file types or sanitize filenames. '
        'Attackers can upload executable files, web shells, or malicious scripts that could '
        'be executed on the server.'
    )
    
    add_heading_custom(doc, '5.6 Information Disclosure (CWE-200)', 2)
    p = doc.add_paragraph()
    p.add_run('Severity: ').bold = True
    p.add_run('MEDIUM')
    p.runs[1].font.color.rgb = RGBColor(255, 165, 0)
    
    doc.add_paragraph(
        'A debug endpoint exposes sensitive system information including the secret key, '
        'environment variables, database path, and Python version. This information aids '
        'attackers in crafting targeted attacks.'
    )
    
    add_heading_custom(doc, '5.7 Path Traversal (CWE-22)', 2)
    p = doc.add_paragraph()
    p.add_run('Severity: ').bold = True
    p.add_run('MEDIUM')
    p.runs[1].font.color.rgb = RGBColor(255, 165, 0)
    
    doc.add_paragraph(
        'File serving functionality does not sanitize the filename parameter, allowing '
        'directory traversal sequences (../) to access files outside the intended '
        'upload directory.'
    )
    
    add_heading_custom(doc, '5.8 Verbose Error Messages (CWE-209)', 2)
    p = doc.add_paragraph()
    p.add_run('Severity: ').bold = True
    p.add_run('LOW')
    p.runs[1].font.color.rgb = RGBColor(255, 200, 0)
    
    doc.add_paragraph(
        'Detailed database error messages are displayed to users, revealing internal '
        'implementation details and potentially aiding attackers in SQL injection attempts.'
    )
    
    add_heading_custom(doc, '5.9 Missing Input Validation (CWE-20)', 2)
    p = doc.add_paragraph()
    p.add_run('Severity: ').bold = True
    p.add_run('MEDIUM')
    p.runs[1].font.color.rgb = RGBColor(255, 165, 0)
    
    doc.add_paragraph(
        'User input is used without validation or sanitization, creating opportunities '
        'for various injection attacks and unexpected application behavior.'
    )
    
    doc.add_page_break()
    
    add_heading_custom(doc, '6. Risk Severity Assessment', 1)
    doc.add_paragraph(
        'The following table summarizes all identified vulnerabilities with their '
        'risk ratings based on CVSS principles and business impact.'
    )
    
    risk_table = doc.add_table(rows=10, cols=4)
    risk_table.style = 'Table Grid'
    
    risk_data = [
        ('Vulnerability', 'CWE', 'Severity', 'Business Impact'),
        ('SQL Injection', 'CWE-89', 'CRITICAL', 'Complete database compromise, data theft'),
        ('Hardcoded Secret', 'CWE-798', 'HIGH', 'Session hijacking, unauthorized access'),
        ('Debug Mode', 'CWE-489', 'HIGH', 'Remote code execution, information leak'),
        ('Plaintext Passwords', 'CWE-256', 'HIGH', 'Credential theft, account takeover'),
        ('Unsafe File Upload', 'CWE-434', 'MEDIUM', 'Malware distribution, server compromise'),
        ('Information Disclosure', 'CWE-200', 'MEDIUM', 'Reconnaissance aid for attackers'),
        ('Path Traversal', 'CWE-22', 'MEDIUM', 'Unauthorized file access'),
        ('Verbose Errors', 'CWE-209', 'LOW', 'Information leakage for reconnaissance'),
        ('Missing Validation', 'CWE-20', 'MEDIUM', 'Multiple attack vectors enabled')
    ]
    
    for i, row_data in enumerate(risk_data):
        row = risk_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                for paragraph in row.cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    add_heading_custom(doc, '7. Remediation Steps', 1)
    doc.add_paragraph(
        'The following remediation measures were implemented in the fixed version of the application:'
    )
    
    remediations = [
        ('Use Parameterized Queries', 
         'Replace all string concatenation in SQL queries with parameterized statements. '
         'This is the most effective defense against SQL injection.'),
        ('Environment-Based Secrets',
         'Load sensitive configuration from environment variables rather than hardcoding. '
         'Use a .env file for development and proper secret management for production.'),
        ('Password Hashing',
         'Implement bcrypt or PBKDF2 password hashing using werkzeug.security. Never store '
         'plaintext passwords.'),
        ('Disable Debug Mode',
         'Ensure debug mode is disabled in all production environments. Set FLASK_ENV=production.'),
        ('File Upload Validation',
         'Validate file extensions against an allowlist, check MIME types, use secure_filename(), '
         'and scan uploads for malware.'),
        ('Remove Debug Endpoints',
         'Eliminate any endpoints that expose system information. Implement proper logging instead.'),
        ('Input Validation',
         'Validate all user input for type, length, format, and range. Reject or sanitize '
         'unexpected input.'),
        ('Error Handling',
         'Display generic error messages to users while logging detailed errors internally.')
    ]
    
    for title, description in remediations:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(description)
    
    add_heading_custom(doc, '8. Secure Coding Best Practices', 1)
    
    best_practices = [
        'Never trust user input - validate and sanitize all data from external sources',
        'Use parameterized queries or ORM to prevent injection attacks',
        'Store passwords using strong adaptive hashing algorithms (bcrypt, Argon2)',
        'Keep secrets out of source code - use environment variables or secret management',
        'Disable debug features and verbose error messages in production',
        'Implement principle of least privilege for database and file system access',
        'Validate and sanitize all file uploads - check extensions, MIME types, and content',
        'Use security headers (CSP, X-Frame-Options, X-Content-Type-Options)',
        'Implement proper session management with secure, httponly, samesite cookies',
        'Regularly update dependencies and scan for known vulnerabilities'
    ]
    
    for practice in best_practices:
        doc.add_paragraph(practice, style='List Bullet')
    
    add_heading_custom(doc, '9. Comparison: Vulnerable vs. Secure Version', 1)
    doc.add_paragraph(
        'The remediation process transformed the vulnerable application into a secure version '
        'through systematic application of secure coding principles. The secure version '
        'maintains all original functionality while eliminating identified security risks.'
    )
    
    comparison_table = doc.add_table(rows=8, cols=3)
    comparison_table.style = 'Table Grid'
    
    comparison_data = [
        ('Security Aspect', 'Vulnerable Version', 'Secure Version'),
        ('SQL Queries', 'String concatenation', 'Parameterized queries'),
        ('Password Storage', 'Plaintext', 'PBKDF2 hashed with salt'),
        ('Secret Management', 'Hardcoded in source', 'Environment variables'),
        ('Debug Mode', 'Enabled', 'Disabled'),
        ('File Uploads', 'No validation', 'Extension and content validation'),
        ('Error Messages', 'Detailed SQL errors', 'Generic user messages'),
        ('Input Handling', 'Direct use', 'Validation and sanitization')
    ]
    
    for i, row_data in enumerate(comparison_data):
        row = comparison_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                for paragraph in row.cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    add_heading_custom(doc, '10. Conclusion', 1)
    doc.add_paragraph(
        'This secure coding review successfully identified and remediated nine security '
        'vulnerabilities in a Flask web application. The exercise demonstrated the importance '
        'of combining automated tools with manual code review for comprehensive security assessment.'
    )
    doc.add_paragraph(
        'The remediation process emphasized that security must be integrated into the development '
        'lifecycle rather than treated as an afterthought. Key lessons include the critical '
        'importance of input validation, secure credential storage, and proper configuration '
        'management.'
    )
    doc.add_paragraph(
        'Moving forward, organizations should implement secure coding training, regular security '
        'assessments, and automated security testing in CI/CD pipelines to maintain application '
        'security posture.'
    )
    
    add_heading_custom(doc, '11. References', 1)
    
    references = [
        'OWASP Foundation. (2021). OWASP Top 10:2021. Retrieved from https://owasp.org/Top10/',
        'MITRE Corporation. (2023). CWE Top 25 Most Dangerous Software Weaknesses. '
        'Retrieved from https://cwe.mitre.org/top25/',
        'Python Software Foundation. (2023). Flask Documentation - Security Considerations. '
        'Retrieved from https://flask.palletsprojects.com/en/2.3.x/security/',
        'Bandit Documentation. (2023). Security Linter for Python. '
        'Retrieved from https://bandit.readthedocs.io/',
        'Semgrep. (2023). Lightweight static analysis for many languages. '
        'Retrieved from https://semgrep.dev/'
    ]
    
    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')
    
    report_path = os.path.join('reports', 'Secure_Coding_Review_Report.docx')
    os.makedirs('reports', exist_ok=True)
    doc.save(report_path)
    
    print(f"Professional report generated: {report_path}")
    return report_path

if __name__ == '__main__':
    generate_report()
